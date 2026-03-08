"""
Generate the Communication Architecture Cost Summary Document (DOCX)
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Georgia'
font.size = Pt(11)
font.color.rgb = RGBColor(0x3F, 0x3E, 0x3C)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.4

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Georgia'
    hs.font.color.rgb = RGBColor(0x3F, 0x3E, 0x3C)
    hs.paragraph_format.space_before = Pt(24 if level == 1 else 18 if level == 2 else 12)
    hs.paragraph_format.space_after = Pt(8)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)

ROSE = RGBColor(0x9C, 0x6F, 0x6E)
OLIVE = RGBColor(0xC7, 0xAE, 0x8C)
DARK = RGBColor(0x3F, 0x3E, 0x3C)

def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_______________________________________________')
    run.font.color.rgb = OLIVE
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(16)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Georgia'

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Georgia'
            if c_idx >= len(headers) - 2 and '$' in str(val):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    return table

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)

def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = ROSE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


# ===== TITLE PAGE =====
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ROSES OS')
run.font.size = Pt(32)
run.font.name = 'Georgia'
run.font.color.rgb = DARK
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Communication Architecture')
run.font.size = Pt(18)
run.font.name = 'Georgia'
run.font.color.rgb = ROSE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Cost Summary')
run.font.size = Pt(18)
run.font.name = 'Georgia'
run.font.color.rgb = ROSE

for _ in range(2):
    doc.add_paragraph()

add_divider(doc)

meta_info = [
    ('Prepared by', 'Jennifer Brooke Lawless'),
    ('Date', 'March 8, 2026'),
    ('Version', '1.0 -- For Team Review'),
    ('Source', 'Communication Architecture for the Age of AI (V2, 2025-2026)'),
]
for label, value in meta_info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Georgia'
    run = p.add_run(value)
    run.font.size = Pt(10)
    run.font.name = 'Georgia'

doc.add_page_break()

# ===== OVERVIEW =====
doc.add_heading('Overview', level=1)

doc.add_paragraph(
    'This document provides a cost summary for implementing the strategic initiatives '
    'outlined in the Communication Architecture for the Age of AI PDF. This is a standalone '
    'estimate covering the GEO (Generative Engine Optimization) strategy, platform ecosystem '
    'buildout, partnership engagement, and supporting infrastructure described in that '
    '21-page strategic framework.'
)

add_quote(doc,
    'This cost summary is separate from and complementary to the Scope of Work & Economic '
    'Proposal (which covers platform development, brand book, manuals, and design system). '
    'Some items overlap -- those are noted where applicable.'
)

add_divider(doc)

# ===== SECTION 1: PRE-LAUNCH =====
doc.add_heading('1. Pre-Launch Strategy Initiatives (5-8 Month Window)', level=1)

doc.add_paragraph(
    'The Communication Architecture identifies a critical 5-8 month pre-launch window '
    'before the book publication. The following workstreams are required to prepare the '
    'knowledge ecosystem.'
)

# 1.1 Substack
doc.add_heading('1.1 Substack Migration & Publication Setup', level=2)
doc.add_paragraph(
    'Migrate 29-issue newsletter archive to Substack. Establish as named, branded publication. '
    'Configure paid subscription tier. Set up recommendation network integrations.'
)
doc.add_paragraph('Deliverables: Branded Substack publication live with full archive, paid tier configured, discovery settings optimized.')
add_table(doc,
    ['', 'Market Value', 'Contribution Rate (10%)'],
    [['Total', '$4,500', '$450']]
)

# 1.2 Blog Essays
doc.add_heading('1.2 Long-Form Blog Essays -- GEO Core Content', level=2)
doc.add_paragraph(
    'Write and publish 8-12 definitional long-form essays on key Roses OS concepts '
    '(Coherence, Remembrance, The Rose Practice, Inner Technology, The Field, Loving Presence). '
    'Structured with proper headers, schema-ready, bilingual (EN + ES). Each 2,000-4,000 words, '
    'structured to directly answer AI-asked questions.'
)
add_table(doc,
    ['', 'Market Value', 'Contribution Rate (10%)'],
    [['Total', '$18,000', '$1,800']]
)

# 1.3 Schema
doc.add_heading('1.3 Schema Markup & Technical Infrastructure', level=2)
doc.add_paragraph(
    'Implement schema.org markup across the Roses OS website (Organization, Person, Course, '
    'Article, FAQPage, DefinedTerm). Optimize page titles, meta descriptions, clean URLs, '
    'alt text, sitemap. Add transcripts infrastructure for audio/video content.'
)
add_table(doc,
    ['', 'Market Value', 'Contribution Rate (10%)'],
    [['Total', '$12,000', '$1,200']]
)
add_note(doc, 'Note: This overlaps with the "GEO / Content Architecture" project in the Scope of Work ($12,000 market). If engaged through that milestone, this line item would be consolidated.')

# 1.4 YouTube Spanish
doc.add_heading('1.4 YouTube Channel -- Spanish Language', level=2)
doc.add_paragraph(
    'Launch dedicated Spanish-language YouTube channel. Establish conceptually structured content '
    'framework (titles, descriptions, chapter markers, transcripts aligned to Roses OS vocabulary). '
    'Initial 8-12 videos produced and optimized.'
)
add_table(doc,
    ['', 'Market Value', 'Contribution Rate (10%)'],
    [['Total', '$22,000', '$2,200']]
)

# 1.5 YouTube English
doc.add_heading('1.5 YouTube Channel -- English Language', level=2)
doc.add_paragraph(
    'Launch English-language YouTube channel with same conceptual structure. '
    'Initial 6-8 videos produced and optimized.'
)
add_table(doc,
    ['', 'Market Value', 'Contribution Rate (10%)'],
    [['Total', '$16,000', '$1,600']]
)

# 1.6 Podcasts
doc.add_heading('1.6 Citation Network -- Strategic Podcast Appearances', level=2)
doc.add_paragraph(
    'Research, outreach, booking, and preparation for 2-3 strategic podcast appearances on '
    'credible shows in consciousness, leadership, and inner development. Priority on '
    'Spanish-language podcasts. Includes pre-appearance coaching, talking point development, '
    'and post-appearance content repurposing.'
)
add_table(doc,
    ['', 'Market Value', 'Contribution Rate (10%)'],
    [['Total', '$9,000', '$900']]
)

# 1.7 Community Voice
doc.add_heading('1.7 Community Voice Program', level=2)
doc.add_paragraph(
    'Design and launch structured program for collecting practitioner essays, written '
    'testimonials, public reflections using Roses OS language. Create submission framework, '
    'editorial guidelines, and publishing workflow.'
)
add_table(doc,
    ['', 'Market Value', 'Contribution Rate (10%)'],
    [['Total', '$6,000', '$600']]
)

# 1.8 Wikipedia
doc.add_heading('1.8 Wikipedia Groundwork', level=2)
doc.add_paragraph(
    'Research Wikipedia notability requirements. Map citation gaps. Develop sourcing strategy. '
    'Prepare draft entries for the school, founder, and key concepts. Accumulate qualifying '
    'independent citations.'
)
add_table(doc,
    ['', 'Market Value', 'Contribution Rate (10%)'],
    [['Total', '$5,000', '$500']]
)

add_divider(doc)

# ===== SECTION 2: CONVERSION BRIDGE =====
doc.add_heading('2. Conversion Bridge Design', level=1)

doc.add_paragraph(
    'The Communication Architecture describes a 5-step conversion journey from AI '
    'recommendation to enrolled student:'
)

steps = [
    'AI recommendation -- the person discovers Roses OS through a response to a question they asked',
    'Structured arrival -- they land on a page that clearly explains the framework',
    'Depth confirmation -- they read an essay, watch a talk, or explore the newsletter',
    'Community signal -- they encounter voices of students and practitioners',
    'Clear invitation -- a specific, well-designed pathway to the next step',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    p.add_run(step)

doc.add_paragraph()
doc.add_paragraph(
    'Scope: Design and build the full conversion bridge across the Roses OS website. '
    'Responsive, bilingual (EN + ES).'
)
add_table(doc,
    ['', 'Market Value', 'Contribution Rate (10%)'],
    [['Total', '$14,000', '$1,400']]
)

add_divider(doc)

# ===== SECTION 3: PARTNERSHIPS =====
doc.add_heading('3. Strategic Partnerships', level=1)

doc.add_paragraph(
    'The Communication Architecture identifies three partnership profiles essential to '
    'accelerating the ecosystem. These represent engagement costs, not ongoing retainers.'
)

doc.add_heading('3.1 Digital Infrastructure Partner', level=2)
doc.add_paragraph(
    'Engage a systematic content distribution partner to build repeatable "build once, '
    'distribute forever" content operations. Establish workflows for turning founder ideas '
    'into structured multi-platform content.'
)
add_table(doc,
    ['', 'Market Value', 'Contribution Rate (10%)'],
    [['Total', '$15,000', '$1,500']]
)

doc.add_heading('3.2 YouTube Strategy Partner', level=2)
doc.add_paragraph(
    'Engage a YouTube thought-leader strategist for channel strategy, algorithmic '
    'optimization, and content structure consulting for both EN + ES channels.'
)
add_table(doc,
    ['', 'Market Value', 'Contribution Rate (10%)'],
    [['Total', '$10,000', '$1,000']]
)

doc.add_heading('3.3 Hispanic Podcast Ecosystem Partner', level=2)
doc.add_paragraph(
    'Identify, outreach, and establish relationship with 1-2 high-reach Spanish-language '
    'podcasters in the consciousness/leadership space. Coordinate appearances timed to '
    'book pre-launch.'
)
add_table(doc,
    ['', 'Market Value', 'Contribution Rate (10%)'],
    [['Total', '$8,000', '$800']]
)

add_divider(doc)

# ===== SECTION 4: MEASUREMENT =====
doc.add_heading('4. Measurement & Monitoring Framework', level=1)

doc.add_paragraph(
    'Set up AI mention monitoring (EN + ES), citation tracking, inbound inquiry quality '
    'tracking, Substack analytics, community voice volume measurement, and enrollment '
    'source analysis.'
)
add_table(doc,
    ['', 'Market Value', 'Contribution Rate (10%)'],
    [['Total', '$6,500', '$650']]
)

add_divider(doc)
doc.add_page_break()

# ===== SECTION 5: SUMMARY =====
doc.add_heading('5. Cost Summary', level=1)

doc.add_heading('By Category', level=2)
add_table(doc,
    ['Category', 'Market Value', 'Contribution Rate'],
    [
        ['Pre-Launch Initiatives (1.1-1.8)', '$92,500', '$9,250'],
        ['Conversion Bridge (2)', '$14,000', '$1,400'],
        ['Strategic Partnerships (3.1-3.3)', '$33,000', '$3,300'],
        ['Measurement Framework (4)', '$6,500', '$650'],
    ]
)

doc.add_heading('Grand Total -- Communication Architecture Implementation', level=2)
add_table(doc,
    ['', 'Market Value', 'Contribution Rate', 'Savings'],
    [
        ['Total', '$146,000', '$14,600', '$131,400 (90%)'],
    ]
)

add_quote(doc,
    '$131,400 in value contributed directly to the mission -- in the same spirit '
    'of conscious contribution that ROSES OS extends to its community.'
)

add_divider(doc)

# ===== SECTION 6: OVERLAP =====
doc.add_heading('6. Overlap with Existing Scope of Work', level=1)

doc.add_paragraph(
    'The following items overlap with projects already defined in the Scope of Work '
    '& Economic Proposal:'
)

add_table(doc,
    ['This Document', 'Scope of Work Equivalent', 'Overlap'],
    [
        ['Schema Markup ($12,000)', 'GEO / Content Architecture ($12,000)', 'Full overlap'],
        ['Conversion Bridge ($14,000)', 'Platform UX Refinement ($15,000)', 'Partial overlap'],
    ]
)

doc.add_paragraph('Adjusted total (removing full overlaps): $134,000 market / $13,400 contribution')
doc.add_paragraph()

doc.add_heading('Combined Project Totals (De-Duplicated)', level=3)
add_table(doc,
    ['Scope', 'Market Value', 'Contribution Rate'],
    [
        ['Scope of Work -- Phases 1 & 2', '$205,000', '$20,500'],
        ['Communication Architecture (de-duplicated)', '$134,000', '$13,400'],
        ['Combined Total', '$339,000', '$33,900'],
    ]
)

add_divider(doc)

# ===== SECTION 7: PRIORITIZATION =====
doc.add_heading('7. Prioritization Recommendation', level=1)

doc.add_paragraph(
    'Based on the Communication Architecture\'s emphasis on the 5-8 month pre-launch window:'
)

doc.add_heading('Immediate (Months 1-2)', level=3)
items_immediate = [
    'Substack Migration -- $4,500 / $450',
    'Schema Markup & Technical Infrastructure -- $12,000 / $1,200',
    'First 4 GEO Blog Essays -- ~$8,000 / $800',
]
for item in items_immediate:
    p = doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Near-Term (Months 2-4)', level=3)
items_near = [
    'Spanish YouTube Channel Launch -- $22,000 / $2,200',
    'Podcast Outreach & First Appearances -- $9,000 / $900',
    'Community Voice Program Launch -- $6,000 / $600',
]
for item in items_near:
    p = doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Pre-Launch (Months 4-8)', level=3)
items_pre = [
    'Remaining GEO Blog Essays -- ~$10,000 / $1,000',
    'English YouTube Channel -- $16,000 / $1,600',
    'Conversion Bridge Design -- $14,000 / $1,400',
    'Strategic Partnerships Engaged -- $33,000 / $3,300',
]
for item in items_pre:
    p = doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Ongoing', level=3)
items_ongoing = [
    'Measurement Framework -- $6,500 / $650',
    'Wikipedia Groundwork -- $5,000 / $500',
]
for item in items_ongoing:
    p = doc.add_paragraph(item, style='List Bullet')

add_divider(doc)

# ===== SECTION 8: NOTES =====
doc.add_heading('8. Notes', level=1)

notes = [
    'All pricing follows the same milestone-based, outcome-driven structure as the Scope of Work -- payment tied to deliverable completion and acceptance, not hours worked.',
    'The 90% contribution rate reflects the same mission-aligned pricing model, with the difference settleable through equity, program enrollment, and/or Numa land allocation.',
    'The Communication Architecture PDF does not assign ownership to specific workstreams. Several initiatives (podcast appearances, community voice, founder visibility) require founder-led participation that cannot be outsourced.',
    'Partnership costs (Section 3) are estimates for initial engagement. Ongoing partnership terms would be negotiated separately.',
    'This is a starting point for conversation. All scope, pricing, and terms are open for discussion.',
]

for note in notes:
    p = doc.add_paragraph(note, style='List Bullet')

add_divider(doc)

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ROSES OS')
run.font.size = Pt(14)
run.font.name = 'Georgia'
run.font.color.rgb = DARK
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Communication Architecture Cost Summary')
run.font.size = Pt(10)
run.font.name = 'Georgia'
run.font.color.rgb = ROSE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('March 2026 -- Version 1.0')
run.font.size = Pt(9)
run.font.name = 'Georgia'
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output_path = os.path.join(os.path.dirname(__file__), 'communication-architecture-cost-summary.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
