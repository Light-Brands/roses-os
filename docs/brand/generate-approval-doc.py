"""
Generate the ROSES OS Brand Book — Founder Approval Document (DOCX)
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Georgia'
font.size = Pt(11)
font.color.rgb = RGBColor(0x3F, 0x3E, 0x3C)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.4

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Georgia'
    hs.font.color.rgb = RGBColor(0x3F, 0x3E, 0x3C)
    hs.paragraph_format.space_before = Pt(24 if level == 1 else 18 if level == 2 else 12)
    hs.paragraph_format.space_after = Pt(8)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)

def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_______________________________________________')
    run.font.color.rgb = RGBColor(0xC7, 0xAE, 0x8C)
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(16)

def add_decision_box(doc, questions):
    """Add a decision/approval box with checkboxes."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('DECISIONS NEEDED')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x9C, 0x6F, 0x6E)
    for q in questions:
        p = doc.add_paragraph()
        run = p.add_run(f'\u2610  {q}')
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x3F, 0x3E, 0x3C)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)

def add_notes_area(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('FOUNDER NOTES:')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x9E, 0x95, 0x6B)
    for _ in range(4):
        p = doc.add_paragraph()
        run = p.add_run('_' * 80)
        run.font.color.rgb = RGBColor(0xC7, 0xAE, 0x8C)
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(8)

def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f'\u201c{text}\u201d')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x9C, 0x6F, 0x6E)


# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ROSES OS')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x9C, 0x6F, 0x6E)
run.font.name = 'Georgia'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Brand Book')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x3F, 0x3E, 0x3C)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Founder Review & Approval Document')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x9E, 0x95, 0x6B)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared for the Guardians of ROSES OS')
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = RGBColor(0xA8, 0x89, 0x6D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('February 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0xA8, 0x89, 0x6D)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Contents', level=1)
doc.add_paragraph()

toc_items = [
    ('How to Use This Document', '3'),
    ('Section 1: Brand Manifest', '4'),
    ('Section 2: Brand Values', '6'),
    ('Section 3: Perpetual Actions', '9'),
    ('Section 4: Brand Drivers', '12'),
    ('Section 5: Brand Perception', '14'),
    ('Section 6: Extended Audience Narrative', '16'),
    ('Section 7: Brand Personality Sliders', '18'),
    ('Reference: Color Palette & Symbolism', ''),
    ('Reference: Typography & Typeface', ''),
    ('Section 8: Butterfly Wing \u2014 Pending Decisions', '20'),
    ('Section 9: Brand Story / Origin Narrative \u2014 Interview Questions', '22'),
]
for title, _ in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# HOW TO USE THIS DOCUMENT
# ============================================================
doc.add_heading('How to Use This Document', level=1)

doc.add_paragraph(
    'This document contains every brand book content section that needs founder review, '
    'approval, or input before design production can begin. It is organized into three types of content:'
)

p = doc.add_paragraph()
run = p.add_run('Drafted sections awaiting approval ')
run.bold = True
p.add_run(
    '\u2014 Brand Manifest, Brand Values, Perpetual Actions, Brand Drivers, Brand Perception, '
    'Audience Narrative, and Personality Sliders. These are written in the ROSES OS voice and '
    'ready for your review. Approve, modify, or redirect.'
)

p = doc.add_paragraph()
run = p.add_run('Pending decisions ')
run.bold = True
p.add_run(
    '\u2014 The Butterfly Wing section contains visual and strategic decisions that need '
    'founder input before the designer can proceed.'
)

p = doc.add_paragraph()
run = p.add_run('Missing content requiring interviews ')
run.bold = True
p.add_run(
    '\u2014 The Brand Story / Origin Narrative can only be written from founder stories. '
    'Interview questions are included and ready to be answered in whatever format feels most alive.'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Each section includes a ')
run.font.size = Pt(11)
r2 = p.add_run('DECISIONS NEEDED')
r2.bold = True
r2.font.color.rgb = RGBColor(0x9C, 0x6F, 0x6E)
p.add_run(' box and a ')
r3 = p.add_run('FOUNDER NOTES')
r3.bold = True
r3.font.color.rgb = RGBColor(0x9E, 0x95, 0x6B)
p.add_run(' area for your responses.')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Note on live site language: ')
run.bold = True
p.add_run(
    'This document reflects the latest copy from rosesos.com as of February 2026. '
    'Key current language: "A Living Consciousness Ecosystem" \u2022 '
    '"Remember Who You Are" \u2022 '
    '"Technologies of remembrance for those ready to live in coherence" \u2022 '
    '"The next revolution is a revolution of consciousness" \u2022 '
    '5,000+ Initiates \u2022 50+ Countries \u2022 30+ Years.'
)

doc.add_page_break()

# ============================================================
# SECTION 1: BRAND MANIFEST
# ============================================================
doc.add_heading('Section 1: Brand Manifest', level=1)
doc.add_paragraph(
    'The Brand Manifest is a 4\u20136 page opening statement that sets the emotional and philosophical '
    'tone for the entire brand book. It is the first thing someone reads. It should feel like a declaration.'
)

doc.add_heading('What We Have', level=2)
doc.add_paragraph('Strong manifesto-quality lines already live on rosesos.com and across brand documents:')

quotes = [
    'Remember Who You Are',
    'The next revolution is a revolution of consciousness.',
    'Technologies of remembrance for those ready to live in coherence.',
    'What if the intelligence you seek is already within you, waiting to be remembered?',
    'When the Rose awakens, Genius awakens.',
    'You are the mountain.',
    'The way is open. Welcome home.',
]
for q in quotes:
    add_quote(doc, q)

doc.add_heading('Proposed Structure', level=2)
doc.add_paragraph(
    'Each page is designed as a full spread \u2014 breathing room, large Cormorant Garamond type, '
    'silence between lines.'
)

spreads = [
    ('Page 1 \u2014 The Readiness',
     '"The intelligence you seek is already within you,\nwaiting to be remembered."'),
    ('Page 2 \u2014 The Technology',
     '"There is a technology older than language. Simpler than strategy.\n'
     'It restores the conditions where clarity becomes natural."'),
    ('Page 3 \u2014 The Revolution',
     '"The next revolution is a revolution of consciousness."'),
    ('Page 4 \u2014 The Awakening',
     '"When the Rose awakens, Genius awakens."'),
    ('Page 5 \u2014 The Welcome',
     '"Remember who you are.\nThe way is open. Welcome home."'),
]
for title, text in spreads:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.5)
    run2 = p2.add_run(text)
    run2.italic = True
    run2.font.color.rgb = RGBColor(0x9C, 0x6F, 0x6E)

doc.add_paragraph()
doc.add_paragraph('Tone direction: Poetry meets purpose. Short sentences. Line breaks as breathing room. '
                   'Every line states what IS \u2014 truths, affirmations, declarations. Speak from the center.')

add_decision_box(doc, [
    'Approve or modify the five-page manifesto structure above',
    'Should the manifesto reference the lineage, or remain purely philosophical?',
    'Should it name the audience ("For seekers, healers, teachers, leaders, creators...") or stay universal?',
    'Is there an additional opening line or truth that should anchor the first page?',
])
add_notes_area(doc)
doc.add_page_break()

# ============================================================
# SECTION 2: BRAND VALUES
# ============================================================
doc.add_heading('Section 2: Brand Values', level=1)
doc.add_paragraph(
    'Core values, each given a full page in the brand book with: the value name, a subtitle, '
    'and a 2\u20134 sentence description of what it means in practice. Written as affirmations '
    'of what ROSES OS stands for.'
)

doc.add_heading('Proposed Values', level=2)

values = [
    ('SOVEREIGN PRESENCE',
     'Being fully here, in this body, in this moment.',
     'The entire system is built on cultivating presence as the ground of everything. '
     'Presence is where clarity lives, where truth is accessible, where genius awakens.'),
    ('RADICAL SIMPLICITY',
     'Choosing the precise over the complex.',
     'Simple, living inner technology. Decentralized simplicity. '
     'The most powerful instruments are the most precise \u2014 and precision is elegant.'),
    ('LIVING COHERENCE',
     'When inside aligns, outside follows.',
     'The central promise. "When coherence returns, life aligns." '
     'This is the core operating principle \u2014 the organizing intelligence of a life well-lived.'),
    ('SACRED INTEGRITY',
     'Honoring the source, the lineage, and the truth of this work.',
     'This work arrives through transmission. It is carried with reverence. '
     'Every teaching holds the frequency of those who came before.'),
    ('QUIET DEVOTION',
     'Returning daily. Practicing fully. Living the work.',
     'Modern devotion. Calm authority. The kind of discipline that comes from love, '
     'the kind of consistency that comes from meaning.'),
    ('HONEST AWARENESS',
     'Seeing what is, clearly and completely.',
     'The Two Core Questions both begin with awareness. Awareness is the first instrument. '
     'Everything else follows from here.'),
]

for name, subtitle, desc in values:
    doc.add_heading(name, level=3)
    p = doc.add_paragraph()
    run = p.add_run(subtitle)
    run.italic = True
    run.font.color.rgb = RGBColor(0x9C, 0x6F, 0x6E)
    doc.add_paragraph(desc)

add_decision_box(doc, [
    'Review and approve or modify the six proposed values',
    'Are there values the founders feel strongly about that aren\u2019t captured here?',
    'Should values reference the practice directly, or stay at the brand philosophy level?',
    'How many values? (Zunya reference uses 6. Fewer may feel more ROSES OS.)',
])
add_notes_area(doc)
doc.add_page_break()

# ============================================================
# SECTION 3: PERPETUAL ACTIONS
# ============================================================
doc.add_heading('Section 3: Perpetual Actions', level=1)
doc.add_paragraph(
    'Active principles that describe how the ROSES OS community and brand behave in the world. '
    'Each is a living commitment, not a static rule. Proposed naming: "Living [noun]" \u2014 '
    'quieter and more precise than the Zunya format of "Active [noun]".'
)

doc.add_heading('Proposed Perpetual Actions', level=2)

actions = [
    ('LIVING REMEMBRANCE',
     'We return. Every practice, every moment of awareness, is a return to what was always true. '
     'Each session is an uncovering. Each day, a homecoming.'),
    ('LIVING COHERENCE',
     'We align. When inner systems are coherent, outer reality responds. '
     'We restore the conditions for clarity, and clarity organizes everything.'),
    ('LIVING PRESENCE',
     'We stay. In the room, in the body, in the moment. '
     'Presence is the ground we stand on. From here, everything becomes possible.'),
    ('LIVING SOVEREIGNTY',
     'We own our energy. We take full responsibility for our inner state. '
     'Freedom begins here \u2014 in the willingness to be whole.'),
    ('LIVING SIMPLICITY',
     'We choose what is essential. Simplicity is depth. Precision is elegance. '
     'The most powerful instruments are the most direct.'),
    ('LIVING TRANSMISSION',
     'We honor the lineage. This work arrived through specific people, in specific moments. '
     'We carry it with precision and reverence, and pass it forward with care.'),
    ('LIVING DEVOTION',
     'We practice daily. The work is quiet. The discipline is steady. The commitment is whole. '
     'Devotion is the rhythm of remembrance.'),
    ('LIVING SERVICE',
     'We give. When coherence is restored, the natural impulse is contribution. '
     'We serve from fullness. We offer from truth.'),
]

for name, desc in actions:
    doc.add_heading(name, level=3)
    doc.add_paragraph(desc)

add_decision_box(doc, [
    'Review and approve or modify the eight proposed actions',
    'Naming convention: "Living [noun]"? Or something different?',
    'How many? (7\u20138 feels right for ROSES OS. Zunya uses 10.)',
    'Should each get its own full page in the book, or paired on spreads?',
])
add_notes_area(doc)
doc.add_page_break()

# ============================================================
# SECTION 4: BRAND DRIVERS
# ============================================================
doc.add_heading('Section 4: Brand Drivers', level=1)
doc.add_paragraph(
    'Macro-level forces that the brand responds to and advances. These are the "why this matters '
    'beyond the individual" statements \u2014 the larger context in which ROSES OS operates.'
)

doc.add_heading('Proposed Drivers', level=2)

drivers = [
    ('INNER TECHNOLOGY FOR THIS MOMENT',
     'Life moves fast. The inner instruments that serve this moment are simple, precise, and alive. '
     'ROSES OS is inner technology designed for the speed, complexity, and depth of modern life.'),
    ('COHERENCE AS LEADERSHIP',
     'Leaders who are coherent inside create coherence around them. When inner symmetry is restored, '
     'leadership becomes natural, creative, and sovereign. The world is ready for this kind of leadership.'),
    ('ACCESSIBLE CONSCIOUSNESS',
     'These teachings have been transmitted through lineage for decades. ROSES OS makes them structured, '
     'available, and accessible \u2014 honoring their depth and sacred nature while opening the door wider.'),
]

for name, desc in drivers:
    doc.add_heading(name, level=3)
    doc.add_paragraph(desc)

add_decision_box(doc, [
    'Review and approve or modify the three proposed drivers',
    'Are there macro-level forces the founders care deeply about that aren\u2019t captured? '
    '(e.g., decentralization, collective coherence, family/parenting, conscious prosperity)',
    'Should drivers be presented visually (one word per page) or with descriptions like above?',
])
add_notes_area(doc)
doc.add_page_break()

# ============================================================
# SECTION 5: BRAND PERCEPTION
# ============================================================
doc.add_heading('Section 5: Brand Perception', level=1)
doc.add_paragraph(
    'A clear statement of how ROSES OS is perceived in the world. Stated entirely in the affirmative '
    '\u2014 who we are, how we feel, what we embody. This becomes a single designed page in the brand book.'
)

doc.add_heading('Draft \u2014 We Are', level=2)

perceptions = [
    'A living consciousness ecosystem',
    'A remembrance ecosystem',
    'Sacred technology for modern life',
    'A lineage-based transmission',
    'An operating system for coherence',
    'A threshold \u2014 a doorway into presence',
    'A temple in digital form',
    'Calm. Precise. Grounding.',
    'A living field of sovereign beings remembering together',
]
for item in perceptions:
    p = doc.add_paragraph()
    run = p.add_run(f'\u2022  {item}')
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Summary line (from rosesos.com): ')
run.bold = True
p.add_run('"Technologies of remembrance for those ready to live in coherence."')

add_decision_box(doc, [
    'Review and refine the identity statements above',
    'Is the live site summary line the right anchor, or is there a stronger one?',
    'Are there identity statements missing that feel essential?',
])
add_notes_area(doc)
doc.add_page_break()

# ============================================================
# SECTION 6: EXTENDED AUDIENCE NARRATIVE
# ============================================================
doc.add_heading('Section 6: Extended Audience Narrative', level=1)
doc.add_paragraph(
    'A full-page prose description of the ideal participant. Not a bullet list \u2014 a paragraph '
    'that someone reads and thinks, "That\u2019s me."'
)

doc.add_heading('Current Audience Language (rosesos.com)', level=2)
doc.add_paragraph('Homepage: Seekers, Healers, Teachers, Leaders, Creators \u2014 "For those who sense there is more"')
doc.add_paragraph('The Rose page: "For those who feel the world accelerating and know the tools they inherited '
                   'were not built for this moment. For founders, creators, parents, healers, and leaders who need '
                   'clarity, intuition, and inner stability to move through complexity."')
doc.add_paragraph('Community page: "A gathering of individuals who have chosen to walk the path of remembrance together."')

doc.add_heading('Draft Narrative', level=2)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
p.paragraph_format.space_before = Pt(8)
run = p.add_run(
    'You have built something real. You lead, you create, you carry responsibility. '
    'And somewhere inside, you sense something deeper \u2014 the intelligence beneath all of it. '
    'The part of you that already knows.'
)
run.italic = True
run.font.color.rgb = RGBColor(0x9C, 0x6F, 0x6E)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
run = p.add_run('You sense there is more.')
run.italic = True
run.font.color.rgb = RGBColor(0x9C, 0x6F, 0x6E)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
run = p.add_run(
    'You are ready for something essential. Something precise. A way home. '
    'A living path back to the intelligence that already lives within you, waiting to be remembered.'
)
run.italic = True
run.font.color.rgb = RGBColor(0x9C, 0x6F, 0x6E)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
run = p.add_run('You are ready to remember.')
run.italic = True
run.font.color.rgb = RGBColor(0x9C, 0x6F, 0x6E)

doc.add_paragraph()

add_decision_box(doc, [
    'Does this resonate with how the founders see their audience?',
    'The homepage uses "Seekers, Healers, Teachers, Leaders, Creators" while /the-rose adds '
    '"Founders, Parents" \u2014 which is the definitive audience set?',
    'Should the narrative address families/parents specifically, or keep it individual?',
    'Any phrases or feelings that are missing from this portrait?',
])
add_notes_area(doc)
doc.add_page_break()

# ============================================================
# SECTION 7: PERSONALITY SLIDERS
# ============================================================
doc.add_heading('Section 7: Brand Personality Sliders', level=1)
doc.add_paragraph(
    'A visual scale showing where the brand sits between two extremes. '
    'These positions become a designed page in the brand book \u2014 a quick, intuitive snapshot '
    'of the brand\u2019s character.'
)

doc.add_heading('Proposed Slider Positions', level=2)

# Create a table for the sliders
table = doc.add_table(rows=9, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for i, text in enumerate(['Spectrum', 'Position', 'Notes']):
    cell = table.rows[0].cells[i]
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

sliders = [
    ('INCLUSIVE \u2194 EXCLUSIVE', 'Center-left', 'Accessible but initiatory'),
    ('LOUD \u2194 QUIET', 'Far right', 'Quiet'),
    ('PLAYFUL \u2194 SERIOUS', 'Right-of-center', 'Serious but not heavy'),
    ('CONFORMIST \u2194 REBELLIOUS', 'Right-of-center', 'Quietly rebellious'),
    ('MINIMAL \u2194 ORNATE', 'Far left', 'Minimal'),
    ('ANCIENT \u2194 MODERN', 'Center', 'Sacred-tech blend'),
    ('PERSONAL \u2194 UNIVERSAL', 'Center-right', 'Universal principles, personal practice'),
    ('PREMIUM \u2194 ORDINARY', 'Left-of-center', 'Premium but accessible'),
]

for i, (spectrum, position, notes) in enumerate(sliders):
    row = table.rows[i + 1]
    row.cells[0].text = spectrum
    row.cells[1].text = position
    row.cells[2].text = notes
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

add_decision_box(doc, [
    'Approve or adjust the slider positions',
    'Add or remove any spectrums?',
    'Does "Quietly rebellious" feel right, or is there a better way to describe the brand\u2019s relationship to convention?',
])
add_notes_area(doc)
doc.add_page_break()

# ============================================================
# REFERENCE: COLOR PALETTE & SYMBOLISM
# ============================================================
doc.add_heading('Reference: Color Palette & Symbolism', level=1)
doc.add_paragraph(
    'The complete ROSES OS color system. Every color carries intention \u2014 these are symbolic decisions '
    'rooted in the brand\u2019s philosophy and lineage. Included here for reference alongside the approval sections.'
)

doc.add_heading('Primary Signature Color', level=2)

# Helper to set cell background color
def set_cell_shading(cell, hex_color):
    """Set cell background shading."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): hex_color,
    })
    shading.append(shading_elem)

# Signature color - large swatch
sig_table = doc.add_table(rows=1, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_swatch = sig_table.rows[0].cells[0]
cell_info = sig_table.rows[0].cells[1]

set_cell_shading(cell_swatch, '9C6F6E')
p = cell_swatch.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')
run.font.size = Pt(24)

p = cell_info.paragraphs[0]
run = p.add_run('ROSE CLAY MAUVE')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x9C, 0x6F, 0x6E)
p2 = cell_info.add_paragraph()
run2 = p2.add_run('#9C6F6E')
run2.font.size = Pt(11)
run2.bold = True
p3 = cell_info.add_paragraph()
run3 = p3.add_run('The human interface layer \u2014 the signature color of ROSES OS.\n'
                    'Divine love made material. The rose-pink of the butterfly wing.\n'
                    'Biological, geological, devotional.')
run3.font.size = Pt(9.5)
run3.font.color.rgb = RGBColor(0x3F, 0x3E, 0x3C)

doc.add_paragraph()
doc.add_heading('Full Palette', level=2)

# Color palette table
colors = [
    ('9C6F6E', 'Rose Clay Mauve', '#9C6F6E', 'Signature',
     'Divine love made material. The Heart Chakra\u2019s highest expression (I Am Love). Grounded in clay and mineral.'),
    ('9B6A66', 'Warm Rose-Clay Brown', '#9B6A66', 'Primary Warm',
     'The body remembering. Terracotta dust, sun-warmed stone. Where spirit enters matter.'),
    ('C4836C', 'Light Terracotta', '#C4836C', 'Primary Warm',
     'Warmth without condition. Sun-warmed clay, soft earth. The feeling of being held by the land.'),
    ('A8896D', 'Gilded Clay', '#A8896D', 'Neutral',
     'Sacred labor. Golden-tan, earthy, sun-kissed. The color of hands that have worked with devotion.'),
    ('C7AE8C', 'Honeyed Stone', '#C7AE8C', 'Neutral',
     'Ancient architecture. The color of temples, sandstone, and structures built to last centuries.'),
    ('EBD6C1', 'Peach Sand', '#EBD6C1', 'Neutral',
     'Soft ground. The color of dawn on sand. Gentleness as a foundation.'),
    ('F5E8E2', 'Golden Ether', '#F5E8E2', 'Neutral',
     'The veil between worlds. Parchment ivory with golden cast. Where the material meets the luminous.'),
    ('FFF8E7', 'Cream Veil', '#FFF8E7', 'Neutral',
     'Silence before the word. Warm ivory, the color of emptiness that holds everything.'),
    ('F7F5F2', 'Aura White', '#F7F5F2', 'Neutral',
     'The field itself. Not sterile white but living white. Sacred space. Breathing room.'),
    ('3F3E3C', 'Soft Charcoal', '#3F3E3C', 'Neutral',
     'Grounded truth. Deep warm earth. The color of certainty without severity. Calm authority.'),
    ('9E956B', 'Antique Olive Brass', '#9E956B', 'Accent',
     'Earned radiance. Like gold in a temple \u2014 present but not ostentatious. Sacred detail.'),
]

color_table = doc.add_table(rows=len(colors) + 1, cols=4)
color_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for i, header in enumerate(['Swatch', 'Color Name / HEX', 'Role', 'Symbolic Meaning']):
    cell = color_table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)
    set_cell_shading(cell, 'EEEEEE')

for row_idx, (hex_val, name, hex_code, role, meaning) in enumerate(colors, 1):
    row = color_table.rows[row_idx]

    # Swatch cell
    swatch_cell = row.cells[0]
    set_cell_shading(swatch_cell, hex_val)
    p = swatch_cell.paragraphs[0]
    run = p.add_run('  ')
    # Use contrasting text color for dark swatches
    if hex_val in ('3F3E3C', '9C6F6E', '9B6A66', 'A8896D', '9E956B'):
        run.font.color.rgb = RGBColor(0xF7, 0xF5, 0xF2)
    else:
        run.font.color.rgb = RGBColor(0x3F, 0x3E, 0x3C)

    # Name + HEX
    name_cell = row.cells[1]
    p = name_cell.paragraphs[0]
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(9)
    p2 = name_cell.add_paragraph()
    run2 = p2.add_run(hex_code)
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x9C, 0x6F, 0x6E)

    # Role
    role_cell = row.cells[2]
    role_cell.text = role
    for paragraph in role_cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)

    # Meaning
    meaning_cell = row.cells[3]
    meaning_cell.text = meaning
    for paragraph in meaning_cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(8.5)
            run.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('The chromatic journey: ')
run.bold = True
run.font.size = Pt(10)
p.add_run('From divine love (rose-pink) through embodiment (clay, terracotta) '
           'into sacred architecture (stone, gold) and into the open field of presence (white, silence). '
           'This is the chromatic journey of remembrance: from heart to body to structure to space.')

doc.add_paragraph()
doc.add_heading('Palette Goals', level=2)
goals = [
    'Earthy, grounded, gender-neutral',
    'Still warm, refined, and modern',
    'Feels good for men + women + non-binary audiences',
    'No bright colors. No saturation. Calm only.',
]
for g in goals:
    p = doc.add_paragraph()
    run = p.add_run(f'\u2022  {g}')
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_page_break()

# ============================================================
# REFERENCE: TYPOGRAPHY & TYPEFACE
# ============================================================
doc.add_heading('Reference: Typography & Typeface', level=1)
doc.add_paragraph(
    'ROSES OS uses two typefaces that embody the sacred-tech bridge \u2014 '
    'one devotional, one modern. Together they create the brand\u2019s visual voice.'
)

doc.add_heading('Cormorant Garamond \u2014 The Sacred Layer', level=2)
doc.add_paragraph('Used for: Hero titles, quote spreads, section headers, manifesto pages')
doc.add_paragraph('Character: Timeless, devotional, poetic \u2014 sacred elegance')
doc.add_paragraph()

# Cormorant examples
examples_cormorant = [
    ('Remember Who You Are', 28, False),
    ('When the Rose Awakens, Genius Awakens', 22, False),
    ('The way is open. Welcome home.', 18, True),
    ('A Living Consciousness Ecosystem', 16, False),
]
for text, size, italic in examples_cormorant:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Garamond'  # Closest available system font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x3F, 0x3E, 0x3C)
    if italic:
        run.italic = True
    # Size label
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f'{size}pt \u2014 Cormorant Garamond{"  Italic" if italic else "  Light"}')
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0xA8, 0x89, 0x6D)

doc.add_paragraph()
add_divider(doc)

doc.add_heading('Inter \u2014 The Modern Layer', level=2)
doc.add_paragraph('Used for: Body text, UI labels, navigation, buttons, descriptions')
doc.add_paragraph('Character: Modern, clean, tech-quiet \u2014 slight tech precision')
doc.add_paragraph()

# Inter examples
examples_inter = [
    ('Technologies of remembrance for those ready to live in coherence.', 13, 'Regular'),
    ('The Rose restores coherence across emotional, mental, and energetic systems \u2014 '
     'making intuition precise and presence sovereign.', 11, 'Regular'),
    ('For those who sense there is more', 11, 'Medium'),
    ('BEGIN YOUR JOURNEY', 10, 'Medium \u2014 Button'),
    ('The Rose  |  Offerings  |  Guardians  |  Community  |  Begin', 10, 'Regular \u2014 Navigation'),
]
for text, size, label in examples_inter:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'  # Closest system font to Inter
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x3F, 0x3E, 0x3C)
    # Label
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'{size}pt \u2014 Inter {label}')
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0xA8, 0x89, 0x6D)

doc.add_paragraph()
add_divider(doc)

doc.add_heading('Typography Rules', level=2)
rules = [
    'Never bold aggressively',
    'Always increase line-height \u2014 text should breathe',
    'No dense blocks \u2014 large spacing, minimal weights',
    'Breath between lines \u2014 silence is part of the design',
    'Cormorant Garamond for sacred moments; Inter for modern clarity',
    'The two fonts bridge devotion and technology',
]
for r in rules:
    p = doc.add_paragraph()
    run = p.add_run(f'\u2022  {r}')
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

doc.add_heading('The Elegance-Tech Bridge', level=2)

# Visual bridge table
bridge_table = doc.add_table(rows=1, cols=3)
bridge_table.alignment = WD_TABLE_ALIGNMENT.CENTER

cell_tech = bridge_table.rows[0].cells[0]
cell_bridge = bridge_table.rows[0].cells[1]
cell_dev = bridge_table.rows[0].cells[2]

set_cell_shading(cell_tech, '3F3E3C')
p = cell_tech.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nTECHNOLOGY\n')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xF7, 0xF5, 0xF2)
p2 = cell_tech.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Inter + spacing\n+ restraint\n')
run2.font.size = Pt(8)
run2.font.color.rgb = RGBColor(0xEB, 0xD6, 0xC1)

set_cell_shading(cell_bridge, '9C6F6E')
p = cell_bridge.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nROSE CLAY MAUVE\n')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xF7, 0xF5, 0xF2)
p2 = cell_bridge.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Human Interface\nLayer\n')
run2.font.size = Pt(8)
run2.font.color.rgb = RGBColor(0xF7, 0xF5, 0xF2)

set_cell_shading(cell_dev, '9E956B')
p = cell_dev.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nDEVOTION\n')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xF7, 0xF5, 0xF2)
p2 = cell_dev.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Cormorant + texture\n+ clay\n')
run2.font.size = Pt(8)
run2.font.color.rgb = RGBColor(0xF7, 0xF5, 0xF2)

doc.add_paragraph()
add_quote(doc, 'Think: Apple-level minimalism inside a temple.')

doc.add_page_break()

# ============================================================
# SECTION 8: BUTTERFLY WING — PENDING DECISIONS
# ============================================================
doc.add_heading('Section 8: The Butterfly Wing \u2014 Pending Decisions', level=1)
doc.add_paragraph(
    'The ROSES OS color palette origin story is anchored by the butterfly wing as the natural source '
    'of the rose-pink tone. Full written content exists for this section (2\u20134 pages in the brand book). '
    'Before the designer can produce these pages, the following decisions need founder input.'
)

doc.add_heading('Context', level=2)
doc.add_paragraph(
    'The rose-pink tone at the heart of the ROSES OS palette came from the butterfly wing. '
    'Certain butterfly species carry this exact mineral-pink: warm, muted, alive. Not a cosmetic pink. '
    'A biological one. A color born from the intersection of light and living structure.'
)
doc.add_paragraph(
    'Butterfly wings produce their color through structural coloration \u2014 microscopic scales that '
    'refract light, revealing beauty through architecture. The color is not applied. It is revealed '
    'when the structure is precise. This is how ROSES OS works: coherence is not added from outside '
    '\u2014 it becomes visible when distortion is removed.'
)
add_quote(doc, 'Pink is the color of divine love. The statement of the green heart is "I Love." '
          'The statement of the pink heart is "I Am Love."')

doc.add_heading('Five Parallels Already Written', level=2)
parallels = [
    ('Structural Revelation', 'The butterfly\u2019s color comes from structure, not surface. '
     'ROSES OS restores inner architecture so coherence becomes naturally visible.'),
    ('Transformation Through Dissolution', 'The chrysalis dissolves the old form entirely. '
     'The Rose dissolves distortion. Both reveal what was always there.'),
    ('Divine Love Made Visible', 'The rose-pink of the butterfly wing is the color of divine love '
     '\u2014 unconditional, quiet, and whole. This is the frequency the brand carries.'),
    ('Precision in Delicacy', 'The butterfly wing is microscopically precise yet appears '
     'effortlessly beautiful. ROSES OS is simple, precise inner technology that feels like warmth and silence.'),
    ('Sovereignty in Lightness', 'The butterfly moves freely, guided by an intelligence it does not '
     'need to understand. When coherence returns, life organizes itself.'),
]
for title, desc in parallels:
    p = doc.add_paragraph()
    run = p.add_run(f'{title} \u2014 ')
    run.bold = True
    run.font.size = Pt(10.5)
    p.add_run(desc)

doc.add_heading('Proposed Book Structure', level=2)
pages = [
    'Page 1: Full-bleed macro image of rose-toned wing scales. Overlay text: "The color of divine love, revealed through structure."',
    'Page 2: Color origin narrative \u2014 how the butterfly wing became the source of the palette.',
    'Page 3: The five parallels between the butterfly wing and ROSES OS.',
    'Page 4: Complete color symbolism table \u2014 every palette color with its symbolic meaning.',
]
for page in pages:
    p = doc.add_paragraph()
    run = p.add_run(f'\u2022  {page}')
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

add_decision_box(doc, [
    'Approve the butterfly wing as a secondary visual element in the brand system (yes / no)',
    'Visual approach: Macro photography or scientific illustration?',
    'Placement in book: Before color swatches (as origin story) or after (as deepening)?',
    'Should butterfly wing texture be integrated into the design system as an optional texture '
    '(alongside linen, paper grain, mineral noise)?',
])
add_notes_area(doc)
doc.add_page_break()

# ============================================================
# SECTION 9: BRAND STORY / ORIGIN NARRATIVE
# ============================================================
doc.add_heading('Section 9: Brand Story / Origin Narrative', level=1)
doc.add_paragraph(
    'The brand story is a 2\u20134 page narrative telling the living story of how ROSES OS came to exist. '
    'Written in the brand voice \u2014 quiet, mythic, grounded. This section cannot be written without '
    'founder stories. The questions below are designed to unlock this narrative.'
)

doc.add_heading('What We Already Have', level=2)
doc.add_paragraph(
    'The lineage path: Lewis S. Bostwick (1960s, California) \u2192 Berkeley Psychic Institute '
    '\u2192 Anastasia Plunk \u2192 Angelina Ataíde \u2192 CELARIS (Brazil) + ROSES OS (International, 2026).'
)
doc.add_paragraph(
    'The name etymology from The Codex: ROSES = from rosa and rhodon \u2014 "essence expressed outwardly, '
    'the invisible made perceivable, Spirit unfolding into form." OS = Operating System \u2014 '
    '"the invisible structure that governs function." Numerology: 11/2 \u2014 the bridge archetype.'
)

doc.add_heading('What\u2019s Still Needed', level=2)
missing = [
    'Angelina Ataíde\u2019s personal story \u2014 how she encountered this work, what she felt, what she carried forward',
    'The "why now" moment \u2014 what made this the right time for ROSES OS to launch internationally',
    'Diego, Dara, and Peggy\u2019s entry \u2014 how each guardian found their way to this work',
    'The growth arc \u2014 the living story behind 30+ years, 5,000+ initiates, 50+ countries',
]
for item in missing:
    p = doc.add_paragraph()
    run = p.add_run(f'\u2022  {item}')
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_heading('Proposed Story Structure', level=2)
structure = [
    ('1. The Lineage', 'Where these teachings came from'),
    ('2. The Carrier', 'Angelina\u2019s story'),
    ('3. The Moment', 'Why this became ROSES OS'),
    ('4. The Name', 'What it means'),
    ('5. The Invitation', 'Where it\u2019s going'),
]
for title, desc in structure:
    p = doc.add_paragraph()
    run = p.add_run(f'{title} \u2014 ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph('Tone direction: Written like a quiet origin myth. Think: "There was a time when..." '
                   '\u2014 grounded, reverent, alive.')

add_divider(doc)

doc.add_heading('Interview Questions for Angelina Ataíde', level=2)
doc.add_paragraph('Guardian of Lineage')
doc.add_paragraph()

angelina_qs = [
    'When did you first encounter this work, and what did it feel like?',
    'Was there a single moment when you knew this was yours to carry \u2014 and what was happening in your life at that time?',
    'Over 30 years you have trained thousands of students, therapists, and teachers across Brazil and Portugal; what kept you going in the early years, before there was a community around this?',
    'What does it mean to you to merge spiritual wisdom with practical techniques \u2014 was that always your intention, or did it reveal itself over time?',
    'The lineage moved from Lewis Bostwick through the Berkeley Psychic Institute, through Anastasia Plunk, to you \u2014 what did you receive from those who came before, and what did you add that is yours alone?',
    'When did you first sense that this work was ready to reach the international stage \u2014 and what made ROSES OS the right vehicle for that expansion, alongside CELARIS continuing in Brazil?',
]
for i, q in enumerate(angelina_qs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {q}')
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(4)

# Notes area for Angelina
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
run = p.add_run('ANGELINA\u2019S RESPONSES:')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x9E, 0x95, 0x6B)
for _ in range(8):
    p = doc.add_paragraph()
    run = p.add_run('_' * 80)
    run.font.color.rgb = RGBColor(0xC7, 0xAE, 0x8C)
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(8)

doc.add_page_break()

doc.add_heading('Interview Questions for All Four Guardians', level=2)
doc.add_paragraph('Can be answered individually or together \u2014 in conversation, voice memo, or writing.')
doc.add_paragraph()

guardian_qs = [
    'What is the moment you remember most clearly from your own journey with this work \u2014 the one that confirmed you were on the right path?',
    'ROSES OS launched in 2026 as the international expression of a lineage carried for decades through CELARIS (formerly Escola da Aura & Sue\u00f1os) in Brazil. What made this the right moment to create an international arm with its own identity and architecture?',
    'The work has reached 5,000+ initiates across 50+ countries. How did that happen \u2014 was it intentional growth, or did people simply find it?',
    'What is the one thing about ROSES OS you wish everyone understood immediately?',
    'Diego, you bridge spirit and structure \u2014 what does that mean in practice, and when did you know that was your role here?',
    'Dara, you hold community and programs \u2014 what does it feel like to hold space for others on this path, and what drew you to this work over a decade ago?',
    'Peggy, you have walked the aura reading path for 20+ years \u2014 what has changed in you over that time, and what has remained the same?',
    'If the brand story is a single arc \u2014 from the 1960s in California to today \u2014 what is that arc about? What is being remembered?',
]
for i, q in enumerate(guardian_qs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {q}')
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(4)

# Notes area for guardians
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
run = p.add_run('GUARDIAN RESPONSES:')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x9E, 0x95, 0x6B)
for _ in range(8):
    p = doc.add_paragraph()
    run = p.add_run('_' * 80)
    run.font.color.rgb = RGBColor(0xC7, 0xAE, 0x8C)
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(8)

doc.add_paragraph()

add_decision_box(doc, [
    'Who will conduct or facilitate these conversations?',
    'Format: Recorded conversation, written responses, or voice memos?',
    'Timeline: When can these be gathered?',
])

doc.add_page_break()

# ============================================================
# CLOSING
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('The way is open.')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x9C, 0x6F, 0x6E)
run.italic = True
run.font.name = 'Georgia'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Welcome home.')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x9C, 0x6F, 0x6E)
run.italic = True
run.font.name = 'Georgia'

# -- Save --
output_path = '/home/user/roses-os/docs/brand/ROSES-OS-Brand-Book-Founder-Approval.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
