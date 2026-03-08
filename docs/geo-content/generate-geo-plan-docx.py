"""
Generate the ROSES OS GEO Content Structure — Implementation Plan (DOCX)

Matches the brand styling from docs/brand/generate-approval-doc.py:
  - Font: Georgia
  - Colors: Rose Clay #9C6F6E, Olive #9E956B, Text #3F3E3C, Gold #C7AE8C
  - Heading sizes: H1=22pt, H2=16pt, H3=13pt

Usage:
    python docs/geo-content/generate-geo-plan-docx.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Brand colors --
ROSE_CLAY = RGBColor(0x9C, 0x6F, 0x6E)
OLIVE = RGBColor(0x9E, 0x95, 0x6B)
TEXT = RGBColor(0x3F, 0x3E, 0x3C)
GOLD = RGBColor(0xC7, 0xAE, 0x8C)
MUTED = RGBColor(0x7A, 0x78, 0x75)

# -- Style setup (identical to brand book generator) --
style = doc.styles['Normal']
font = style.font
font.name = 'Georgia'
font.size = Pt(11)
font.color.rgb = TEXT
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.4

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Georgia'
    hs.font.color.rgb = TEXT
    hs.paragraph_format.space_before = Pt(24 if level == 1 else 18 if level == 2 else 12)
    hs.paragraph_format.space_after = Pt(8)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)


# -- Helper functions --

def add_divider():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_______________________________________________')
    run.font.color.rgb = GOLD
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(16)


def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT
    return p


def add_body_bold(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT
    run.bold = True
    return p


def add_bullet(text, indent=0.25):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'•  {text}')
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT
    return p


def add_label(text):
    """Small uppercase label in rose clay."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = ROSE_CLAY
    return p


def add_status_line(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED
    run = p.add_run(value)
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT
    return p


def add_table(headers, rows):
    """Add a styled table with brand colors."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = ROSE_CLAY
        run.font.name = 'Georgia'
        # Light rose background for header
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): 'F5EFED',
        })
        shading.append(shading_elem)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = TEXT
            run.font.name = 'Georgia'

    # Set column widths if possible
    doc.add_paragraph()  # spacing after table
    return table


def add_checklist(items, checked=False):
    """Add checklist items."""
    mark = '\u2611' if checked else '\u2610'
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'{mark}  {item}')
        run.font.size = Pt(10.5)
        run.font.color.rgb = TEXT


# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ROSES OS')
run.font.size = Pt(36)
run.font.color.rgb = ROSE_CLAY
run.font.name = 'Georgia'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GEO Content Structure')
run.font.size = Pt(24)
run.font.color.rgb = TEXT
run.font.name = 'Georgia'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
run = p.add_run('Implementation Plan')
run.font.size = Pt(18)
run.font.color.rgb = MUTED
run.font.name = 'Georgia'
run.italic = True

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('_______________________________________________')
run.font.color.rgb = GOLD
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
run = p.add_run('Phase 1 Complete  |  Phase 2 In Progress')
run.font.size = Pt(11)
run.font.color.rgb = OLIVE
run.font.name = 'Georgia'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
run = p.add_run('March 2026')
run.font.size = Pt(11)
run.font.color.rgb = MUTED
run.font.name = 'Georgia'

doc.add_page_break()

# ============================================================
# CONTEXT
# ============================================================
doc.add_heading('Context', level=1)

add_body(
    'ROSES OS currently has zero geo-targeted or question-answering content. '
    'All pages are static brand/marketing pages. People searching '
    '"meditation classes New York", "aura reading course online", or '
    '"what is rose meditation" will never find the site.'
)

add_body(
    'This plan adds programmatic SEO infrastructure: geo-targeted landing pages '
    'per city, FAQ content with schema markup, a sitemap, robots file, and Course '
    'structured data — all designed to capture organic search traffic and funnel '
    'it into the enrollment flow.'
)

add_divider()

# ============================================================
# WHY THIS MATTERS (SEO STRATEGY)
# ============================================================
doc.add_heading('Why This Matters (SEO Strategy)', level=1)

add_table(
    ['SEO Gap', 'Solution'],
    [
        ['No pages targeting "[service] + [city]" queries', 'Geo landing pages at /meditation/[location]'],
        ['No content answering "what is..." queries', 'FAQ sections with FAQPage schema (Google "People Also Ask")'],
        ['No sitemap.xml', 'Auto-generated sitemap via Next.js'],
        ['No robots.txt', 'Proper robots file blocking admin/teaching'],
        ['No Course structured data', 'Rich results for course dates, prices, format'],
        ['No breadcrumb schema', 'Better SERP display with navigation trail'],
        ['Unused schema generators in seo.tsx', 'Activate generateFAQSchema(), generateBreadcrumbSchema(), etc.'],
        ['Homepage has no metadata', 'Added in Phase 2 — full SEO metadata export'],
    ],
)

add_label('Target queries this captures')
add_bullet('"meditation classes [city]" — geo pages')
add_bullet('"aura reading course online" — /meditation/online page')
add_bullet('"what is rose meditation" — FAQ schema → People Also Ask')
add_bullet('"chakra cleansing meditation" — FAQ content')
add_bullet('"clairvoyance training online" — FAQ + online page')
add_bullet('"energy healing [city]" — geo pages with city keywords')

add_divider()

# ============================================================
# WHAT WE'RE BUILDING
# ============================================================
doc.add_heading('What We\'re Building', level=1)

# -- 1. Geo Landing Pages --
doc.add_heading('1. Geo Landing Pages (20 pages)', level=2)

add_body(
    'Dynamic route at /meditation/[location] — one page per target city '
    'plus an "online" catch-all.'
)

add_label('Cities by timezone (mapped to existing schedule data)')

add_table(
    ['Timezone', 'Cities', 'Schedule Column'],
    [
        ['Pacific Time', 'San Francisco, Los Angeles', 'sanJose'],
        ['Colombia / Central', 'Bogotá, Mexico City', 'bogota'],
        ['Eastern Time', 'New York, Miami, Toronto', 'newYork'],
        ['Brasília Time', 'São Paulo, Rio de Janeiro, Buenos Aires', 'brasilia'],
        ['GMT', 'London, Lisbon', 'london'],
        ['CET', 'Madrid, Barcelona, Paris, Berlin, Paphos', 'madrid'],
        ['Non-geo', 'Online (catch-all)', 'All timezones'],
    ],
)

add_label('Each page includes')
add_bullet('Hero with city-specific title (e.g. "Rose Meditation & Aura Reading in New York")')
add_bullet('Unique intro paragraph per city (prevents duplicate content)')
add_bullet('Course cards (Rose Meditation + Aura Reading Level 1)')
add_bullet('Schedule filtered to the local timezone')
add_bullet('FAQ accordion (15 questions with FAQPage schema)')
add_bullet('CTA linking to enrollment')

add_label('Each page has structured data')
add_bullet('Course schema (schema.org/Course) with real dates and prices')
add_bullet('FAQPage schema for Google "People Also Ask"')
add_bullet('Breadcrumb schema (Home > Meditation > City)')
add_bullet('WebPage schema')

# -- 2. FAQ Content --
doc.add_heading('2. FAQ Content (15 questions)', level=2)

add_body('Questions targeting real search queries, embedded in every geo page:')

add_table(
    ['Category', 'Questions'],
    [
        ['About', 'What is Rose Meditation? · How does aura reading work? · What is chakra cleansing meditation? · What is clairvoyance training? · Is this a religious practice?'],
        ['Logistics', 'Do I need prior experience? · How are classes delivered? · What timezones do you support? · How long is Rose Meditation? · How long is Aura Reading? · Can I take Rose Meditation standalone?'],
        ['Practice', 'What will I learn? · What happens after the course?'],
        ['Pricing', 'How much do courses cost? · What is "pay what feels right"?'],
    ],
)

# -- 3. Sitemap --
doc.add_heading('3. Sitemap (/sitemap.xml)', level=2)
add_body(
    'Auto-generated sitemap including all static routes (homepage, offerings, '
    'community, etc.) and all 20 geo pages. Updates automatically when cities are added.'
)

# -- 4. Robots --
doc.add_heading('4. Robots (/robots.txt)', level=2)
add_body(
    'Allows crawling of all public pages. Blocks /admin/, /api/, /teaching/ '
    'from search engines. Points to sitemap.'
)

# -- 5. Course Schema --
doc.add_heading('5. Course Schema', level=2)
add_body('schema.org/Course structured data for each program offering:')
add_bullet('Course name, description, provider')
add_bullet('CourseInstance with start/end dates, online delivery mode')
add_bullet('Offer with price (starting tier) and availability')

add_divider()

# ============================================================
# TECHNICAL IMPLEMENTATION
# ============================================================
doc.add_heading('Technical Implementation', level=1)

doc.add_heading('Files Created / Modified', level=2)

add_table(
    ['Action', 'File', 'Purpose'],
    [
        ['Edit', 'src/lib/data/types.ts', 'Add GeoLocation and FAQItem interfaces'],
        ['Create', 'src/lib/data/geo-data.ts', '20 locations, 15 FAQs, helper functions'],
        ['Edit', 'src/lib/data/index.ts', 'Re-export geo data'],
        ['Edit', 'src/lib/seo.tsx', 'Add generateCourseSchema()'],
        ['Create', 'src/components/sections/LocalSchedule.tsx', 'Timezone-filtered schedule table'],
        ['Create', 'src/components/sections/FAQSection.tsx', 'Accordion FAQ component'],
        ['Create', 'src/app/(site)/meditation/[location]/page.tsx', 'Dynamic geo landing page'],
        ['Create', 'src/app/sitemap.ts', 'Auto-generated sitemap'],
        ['Create', 'src/app/robots.ts', 'Crawl directives'],
        ['Create', 'docs/geo-content/city-targeting-questions.md', 'Founder review doc for city list'],
    ],
)

doc.add_heading('Architecture Decisions', level=2)

add_bullet('Server-rendered pages — Geo pages are server components for maximum SEO benefit (full HTML delivered to crawlers)')
add_bullet('Static generation — All geo pages are pre-built at deploy time via generateStaticParams() (fast load times)')
add_bullet('Data-driven — All content lives in geo-data.ts following the existing mock-data.ts pattern. Adding a city = adding one data entry')
add_bullet('Reuses existing components — PageHero, ProgramCard, existing SEO utilities')
add_bullet('Unique content per page — Each city gets a unique intro, timezone-specific schedule, and localized keywords to avoid duplicate content penalties')

doc.add_heading('How to Add a New City', level=2)
add_body('Add an entry to src/lib/data/geo-data.ts:')

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.25)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
code_text = (
    "{\n"
    "  id: 'seattle',\n"
    "  slug: 'seattle',\n"
    "  city: 'Seattle',\n"
    "  region: 'Washington',\n"
    "  country: 'United States',\n"
    "  timezoneKey: 'sanJose',\n"
    "  timezoneLabel: 'Pacific Time (PT)',\n"
    "  intro: 'Join Rose Meditation and Aura Reading courses from Seattle...',\n"
    "  keywords: ['meditation classes seattle', 'aura reading course seattle'],\n"
    "}"
)
run = p.add_run(code_text)
run.font.size = Pt(9)
run.font.name = 'Consolas'
run.font.color.rgb = MUTED

add_body('The page at /meditation/seattle is automatically generated with all content, schema, and sitemap inclusion.')

add_divider()

# ============================================================
# PHASE 2: HOMEPAGE METADATA & DOCS
# ============================================================
doc.add_heading('Phase 2: Homepage Metadata & Documentation', level=1)

add_label('Status')
add_body('In Progress')

doc.add_heading('What Was Missing', level=2)
add_body(
    'The homepage (src/app/(site)/page.tsx) — the most important page for SEO — '
    'had no metadata export. Every other page in the site (offerings, the-rose, '
    'community, contact, guardians) already followed the server-wrapper pattern '
    'with metadata, but the homepage was a single \'use client\' file with no '
    'SEO metadata at all.'
)

doc.add_heading('What We Did', level=2)

add_bullet('Renamed src/app/(site)/page.tsx → src/app/(site)/HomeClient.tsx (preserving the existing client component)')
add_bullet('Created new src/app/(site)/page.tsx as a server wrapper with full metadata')
add_bullet('Title: "ROSES OS — Remember Who You Are"')
add_bullet('Description targeting core search queries (rose meditation, aura reading, consciousness)')
add_bullet('12 SEO keywords')
add_bullet('OpenGraph metadata for social sharing')
add_bullet('Updated scope of work to reflect GEO Phase 1 completion')
add_bullet('Generated this implementation plan as a .docx document')

doc.add_heading('Files Modified (Phase 2)', level=2)

add_table(
    ['Action', 'File', 'Purpose'],
    [
        ['Rename', 'page.tsx → HomeClient.tsx', 'Move client component'],
        ['Create', 'src/app/(site)/page.tsx', 'Server wrapper with homepage metadata'],
        ['Edit', 'docs/brand/scope-of-work-and-proposal.md', 'Update GEO milestone status'],
        ['Edit', 'docs/geo-content/implementation-plan.md', 'Add Phase 2 section'],
        ['Create', 'docs/geo-content/generate-geo-plan-docx.py', 'Python script to generate .docx'],
        ['Create', 'docs/geo-content/ROSES-OS-GEO-Content-Plan.docx', 'Generated .docx deliverable'],
    ],
)

add_divider()

# ============================================================
# VERIFICATION CHECKLIST
# ============================================================
doc.add_heading('Verification Checklist', level=1)

doc.add_heading('Phase 1', level=2)
add_checklist([
    'pnpm build — All geo pages generate without errors',
    '/sitemap.xml includes all static + geo routes',
    '/robots.txt returns correct crawl directives',
    'JSON-LD validates on geo pages (Course, FAQ, Breadcrumb schemas)',
    'Geo pages show correct timezone in schedule',
    'Each geo page has unique title, description, and intro content',
    'FAQ accordion opens/closes correctly',
    'CTA links work (enrollment, offerings)',
], checked=True)

doc.add_heading('Phase 2', level=2)
add_checklist([
    'Homepage exports metadata (title, description, keywords, openGraph)',
    'HomeClient.tsx renders correctly when imported by new page.tsx',
    '.docx document generated and matches brand styling',
    'Scope of work updated with GEO progress',
], checked=False)

# ============================================================
# SAVE
# ============================================================
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, 'ROSES-OS-GEO-Content-Plan.docx')
doc.save(output_path)
print(f'Generated: {output_path}')
